from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # Style definitions
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Title Section
    title = doc.add_heading('Project Report: FakeN TruthSeeker', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('A Hybrid AI Framework for Fake News Detection using ML, RAG, and LLMs').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n')

    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "FakeN TruthSeeker is a sophisticated 'Compound AI' system designed to combat the spread of viral misinformation on social media, "
        "particularly targeting the Indian context (WhatsApp forwards). Unlike traditional single-model chatbots, FakeN utilizes a triple-layer "
        "verification architecture: stylistic pattern recognition (ML), factual retrieval-augmented generation (RAG), and high-level cognitive "
        "reasoning (LLM). This synergy ensures high accuracy, significantly reduced hallucinations, and transparency in its verdicts."
    )

    # 2. Problem Statement
    doc.add_heading('2. Problem Statement', level=1)
    doc.add_paragraph(
        "In the current digital age, misinformation spreads faster than facts. Typical challenges include:"
    )
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('The "WhatsApp Forward" Epidemic: ').bold = True
    p.add_run('Emotional and manipulative text patterns that bypass human critical thinking.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('LLM Hallucinations: ').bold = True
    p.add_run('General-purpose AI models often provide out-of-date or entirely fabricated information when asked about local events.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Static Knowledge: ').bold = True
    p.add_run('AI models are frozen at their training cutoff, making them ineffective against today\'s breaking hoaxes.')

    # 3. System Architecture
    doc.add_heading('3. System Architecture', level=1)
    doc.add_paragraph(
        "The system operates through a structured pipeline to process and verify incoming text."
    )
    
    doc.add_heading('3.1 Pattern Recognition (ML Layer)', level=2)
    doc.add_paragraph(
        "A Scikit-Learn based machine learning model analyzes the 'style' of the news. Using TF-IDF vectorization and Logistic Regression, "
        "it identifies characteristics such as excessive emojis, urgent commands ('Forward now!'), and emotional hyperbole. "
        "It generates a 'Suspicion Score' (0-1)."
    )
    
    doc.add_heading('3.2 RAG Engine (Factual Layer)', level=2)
    doc.add_paragraph(
        "The Retrieval-Augmented Generation (RAG) engine acts as the system's librarian. It searches a localized Knowledge Base "
        "containing over 50 verified facts and common hoaxes. This layer grounds the AI in reality, providing specific evidence to support or debunk a claim."
    )
    
    doc.add_heading('3.3 Cognitive Reasoner (Gemini 2.5 Flash)', level=2)
    doc.add_paragraph(
        "The final layer uses the Google Gemini 2.5 Flash model. It acts as a judge, synthesizing the forensic style score and the "
        "retrieved facts to produce a nuanced, human-readable verdict and explanation."
    )

    # 4. Key Features
    doc.add_heading('4.4 Key Features', level=1)
    features = [
        "Skeptical Logic: Identifies when a topic is real but the specific claim is exaggerated.",
        "Transparent Reporting: Clearly displays both pattern scores and factual matches.",
        "Zero-Hallucination Guard: Prioritizes verified data from the local repository over general AI training.",
        "Responsive UI: Glassmorphic chat interface with color-coded risk indicators."
    ]
    for feat in features:
        doc.add_paragraph(feat, style='List Bullet')

    # 5. Implementation Stack
    doc.add_heading('5. Implementation Stack', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology Used'
    
    items = [
        ('Frontend', 'React + Vite (Modern JavaScript)'),
        ('Backend', 'FastAPI (Python)'),
        ('AI Model', 'Google Gemini 2.5 Flash'),
        ('ML Library', 'Scikit-Learn'),
        ('Data Format', 'JSON-based Knowledge Base')
    ]
    for comp, tech in items:
        row_cells = table.add_row().cells
        row_cells[0].text = comp
        row_cells[1].text = tech

    # 6. Conclusion
    doc.add_heading('6. Conclusion', level=1)
    doc.add_paragraph(
        "FakeN TruthSeeker demonstrates that the future of information integrity lies not in single AI models, but in 'Compound AI' systems. "
        "By combining traditional ML forensics with modern RAG and Large Language Models, we have created a platform that is accurate, "
        "up-to-date, and highly resistant to the common pitfalls of modern AI."
    )

    doc.save('FakeN_Project_Report.docx')
    print("Report created successfully!")

if __name__ == "__main__":
    create_report()
